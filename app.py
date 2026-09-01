import streamlit as st
import json
import re
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from groq import Groq
except Exception:
    Groq = None

st.set_page_config(page_title="Generador de Planes Curriculares", layout="wide")

def get_groq_client():
    api_key = st.secrets.get("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("Falta la clave de API en los secretos.")
    if Groq is None:
        raise RuntimeError("No se pudo importar `groq`.")
    return Groq(api_key=api_key)

def extract_json_from_text(text):
    try:
        return json.loads(text)
    except Exception:
        pass
    regex_obj = re.compile(r"(\{.*\})", re.DOTALL)
    for regex in [regex_obj]:
        m = regex.search(text)
        if m:
            try:
                return json.loads(m.group(1))
            except Exception:
                continue
    raise ValueError("No se pudo extraer JSON válido de la respuesta de IA.")

def set_cell_bg(cell, color_hex="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def insert_markdown_text(cell, text, append=False, align_justify=False):
    if not text:
        return
    
    if not append:
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
            
    paragraphs = text.split('\n')
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        p = cell.add_paragraph()
        if align_justify:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        tokens = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', para_text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('***') and token.endswith('***'):
                run = p.add_run(token[3:-3])
                run.bold = True
                run.italic = True
            elif token.startswith('**') and token.endswith('**'):
                run = p.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*'):
                run = p.add_run(token[1:-1])
                run.italic = True
            else:
                run = p.add_run(token)
            
            run.font.name = 'Arial'
            run.font.size = Pt(9)

def build_prompt(tema, semanas, nivel, area, puntos_clave):
    weeks = int(semanas)
    instruccion_extra = f"Puntos clave solicitados por el docente: {puntos_clave}" if puntos_clave else ""
    
    if nivel == "PRIMARIA":
        reglas_nivel = """
        - Contenidos: Escribe "***Semana X***" (en negrita y cursiva), salto de línea, TÍTULO EN MAYÚSCULAS, salto de línea, y viñetas (•) detalladas con los conceptos clave de la semana.
        - Recursos: Repite recursos lógicos y consistentes acordes a primaria (ej. Cuadernos de apuntes, materiales escolares y de apoyo didáctico de la región).
        """
    else:
        reglas_nivel = """
        - Contenidos: Formato técnico avanzado. Escribe "**Semana X**", luego el tema y los conceptos técnicos correspondientes.
        """

    prompt = f"""
Actúa como un pedagogo boliviano experto en diseño curricular del Ministerio de Educación. 
Debes redactar un Plan de Desarrollo Curricular (PDC) EXTREMADAMENTE DETALLADO, EXTENSO Y PEDAGÓGICAMENTE RICO. 
ESTÁ ESTRICTAMENTE PROHIBIDO GENERAR TEXTOS CORTOS, RESÚMENES O FRASES ROBÓTICAS.

Tema a desarrollar: "{tema}"
Área: "{area}"
Nivel educativo: "{nivel}"
Duración: {weeks} semanas.
Contexto regional: Municipio de Ixiamas, Norte de La Paz. {instruccion_extra}

REGLAS ESTRICTAS DE CONTENIDO (CRÍTICAS PARA APROBAR EL PLAN):
1. Objetivo holístico de nivel: Debe ser de almenos 4 lineas (mínimo 35 palabras).
2. Objetivo de aprendizaje: Párrafo amplio , claro y entendible pero no redundante (mínimo 35 palabras). Usa **negrita** para resaltar verbos de acción y conectores (ej. **Fortalecemos**, **asumiendo**, **a través del análisis**, **para promover**). 
3. Momentos del proceso formativo (CRÍTICO): 
   - ¡PROHIBIDO empezar todas las oraciones con "Los estudiantes..."! Usa una redacción fluida, en primera persona del plural (ej: Iniciamos con..., dialogamos sobre..., construimos...).
   - Cada semana debe tener al menos 80 a 100 palabras en esta sección.
   - Describe la **(Práctica)** partiendo desde la experiencia vivida o el contacto directo con la realidad, de forma muy descriptiva.
   - Describe la **(Teoría)** como un análisis profundo, conceptualización y comprensión exhaustiva del tema.
   - Describe la **(Producción)** especificando exactamente qué producto tangible, creativo o intelectual se va a elaborar.
   - Describe la **(Valoración)** con una reflexión ética y comunitaria.
4. Criterios de Evaluación: SOLO debes generar criterios para SER, SABER y HACER. NO generes criterios para decidir.
{reglas_nivel}

Devuelve ÚNICAMENTE un objeto JSON válido con esta estructura exacta:
{{
  "objetivo_holistico_nivel": "Párrafo muy extenso, profundo y articulado del objetivo de nivel (solo ser, saber, hacer)...",
  "objetivo_aprendizaje": "Párrafo amplio y detallado con las **palabras clave** en negrita explicando el qué, cómo y para qué...",
  "criterios_evaluacion": "**SER:**\\n- Detalle amplio 1\\n- Detalle amplio 2\\n**SABER:**\\n- Detalle amplio 1\\n- Detalle amplio 2\\n**HACER:**\\n- Detalle amplio 1\\n- Detalle amplio 2",
  "semanas": [
    {{
      "semana": 1,
      "contenidos": "...",
      "momentos": "Iniciamos la sesión con un diálogo participativo sobre... para conectar con sus vivencias **(Práctica)**. Posteriormente, analizamos a profundidad los conceptos de... comprendiendo su funcionamiento en el entorno **(Teoría)**. Con estos saberes, elaboramos creativamente un... demostrando destreza técnica **(Producción)**. Finalmente, reflexionamos de manera comunitaria sobre la importancia de... para el bienestar de la región **(Valoración)**.",
      "recursos": "...",
      "periodos": "2"
    }}
  ],
  "adaptaciones_curriculares": "Párrafo de 5 lineas detallado explicando estrategias específicas para estudiantes con dificultades de aprendizaje o talentos extraordinarios."
}}
Genera exactamente {weeks} objetos dentro del arreglo "semanas".
"""
    return prompt

def generar_docx(datos_formulario, plan_json):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    rPr = font.element.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'es-ES')
    rPr.append(lang)

    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.0)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if datos_formulario.get("nivel") == "PRIMARIA":
        titulo_principal = "PLAN DE DESARROLLO CURRICULAR PARA EDUCACIÓN PRIMARIA COMUNITARIA VOCACIONAL\n"
    else:
        titulo_principal = "PLAN DE DESARROLLO CURRICULAR PARA EDUCACIÓN SECUNDARIA COMUNITARIA PRODUCTIVA\n"
        
    run_t1 = p_titulo.add_run(titulo_principal)
    run_t1.bold = True
    run_t1.font.size = Pt(11)
    
    run_t2 = p_titulo.add_run(f"PLAN DE DESARROLLO CURRICULAR Nº {datos_formulario.get('nro_plan', '1')}")
    run_t2.bold = True
    run_t2.font.size = Pt(11)

    p_datos = doc.add_paragraph()
    p_datos.add_run("1.  Datos referenciales").bold = True

    t_datos = doc.add_table(rows=5, cols=4)
    t_datos.style = 'Table Grid'
    t_datos.autofit = False
    
    for row in t_datos.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(3.0)
        row.cells[2].width = Inches(1.5)
        row.cells[3].width = Inches(3.0)

    def set_cell_bold(cell, text):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    set_cell_bold(t_datos.cell(0,0), "Distrito educativo")
    t_datos.cell(0,1).text = "Ixiamas"
    set_cell_bold(t_datos.cell(0,2), "Unidad Educativa")
    t_datos.cell(0,3).text = datos_formulario.get("unidad", "")

    set_cell_bold(t_datos.cell(1,0), "Nivel")
    nivel_texto = "Primaria Comunitaria Vocacional" if datos_formulario.get("nivel") == "PRIMARIA" else "Secundaria Comunitaria Productiva"
    t_datos.cell(1,1).text = nivel_texto
    set_cell_bold(t_datos.cell(1,2), "Año de escolaridad/Paralelo")
    t_datos.cell(1,3).text = datos_formulario.get("ano", "")

    set_cell_bold(t_datos.cell(2,0), "Maestra/o")
    t_datos.cell(2,1).text = datos_formulario.get("docente", "")
    t_datos.cell(2,1).merge(t_datos.cell(2,3))

    set_cell_bold(t_datos.cell(3,0), "Área")
    t_datos.cell(3,1).text = datos_formulario.get("area", "") 
    t_datos.cell(3,1).merge(t_datos.cell(3,3))

    set_cell_bold(t_datos.cell(4,0), "Trimestre")
    t_datos.cell(4,1).text = datos_formulario.get("trimestre", "")
    t_datos.cell(4,1).merge(t_datos.cell(4,3))
    
    row_fecha = t_datos.add_row()
    set_cell_bold(row_fecha.cells[0], "Fecha")
    row_fecha.cells[1].text = datos_formulario.get("fechas", "")
    row_fecha.cells[1].merge(row_fecha.cells[3])

    doc.add_paragraph("") 

    p_des = doc.add_paragraph()
    p_des.add_run("2.  Desarrollo\n").bold = True
    p_des.add_run("Objetivo holístico de nivel\n").bold = True
    p_des.add_run(plan_json.get("objetivo_holistico_nivel", ""))

    semanas = plan_json.get("semanas", [])
    num_semanas = len(semanas)
    
    total_filas = num_semanas + 2
    t_plan = doc.add_table(rows=total_filas, cols=6)
    t_plan.style = 'Table Grid'
    t_plan.autofit = False

    anchos = [Inches(1.4), Inches(1.7), Inches(2.5), Inches(1.5), Inches(0.6), Inches(2.3)]

    cabeceras = ["Objetivo de aprendizaje", "Contenidos", "Momentos del proceso formativo", "Recursos", "Períodos", "Criterios de evaluación"]
    for i, txt in enumerate(cabeceras):
        celda = t_plan.cell(0,i)
        celda.width = anchos[i]
        set_cell_bg(celda, "D9E2F3")
        p = celda.paragraphs[0]
        p.add_run(txt).bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for idx, s in enumerate(semanas):
        fila_actual = idx + 1
        for col_idx in range(6):
            t_plan.cell(fila_actual, col_idx).width = anchos[col_idx]
            
        insert_markdown_text(t_plan.cell(fila_actual, 1), s.get("contenidos", ""))
        insert_markdown_text(t_plan.cell(fila_actual, 2), s.get("momentos", ""), align_justify=True)
        insert_markdown_text(t_plan.cell(fila_actual, 3), s.get("recursos", ""))
        
        p_per = t_plan.cell(fila_actual, 4).paragraphs[0] if t_plan.cell(fila_actual, 4).paragraphs else t_plan.cell(fila_actual, 4).add_paragraph()
        p_per.add_run(str(s.get("periodos", ""))).font.size = Pt(9)
        p_per.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    insert_markdown_text(t_plan.cell(1, 0), plan_json.get("objetivo_aprendizaje", ""), align_justify=True)
    insert_markdown_text(t_plan.cell(1, 5), plan_json.get("criterios_evaluacion", ""))

    fila_adapt = num_semanas + 1
    for col_idx in range(6):
        t_plan.cell(fila_adapt, col_idx).width = anchos[col_idx]
        
    cell_adapt_mid = t_plan.cell(fila_adapt, 1)
    p_adapt = cell_adapt_mid.paragraphs[0]
    run_ad = p_adapt.add_run("Adaptaciones curriculares:\n")
    run_ad.bold = True
    run_ad.font.size = Pt(9)
    insert_markdown_text(cell_adapt_mid, plan_json.get("adaptaciones_curriculares", ""), append=True, align_justify=True)

    t_plan.cell(fila_adapt, 1).merge(t_plan.cell(fila_adapt, 4))
    t_plan.cell(1, 0).merge(t_plan.cell(fila_adapt, 0))
    t_plan.cell(1, 5).merge(t_plan.cell(fila_adapt, 5))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

st.title("Generador de Planes Curriculares")

MESES_ESPANOL = {1:"enero", 2:"febrero", 3:"marzo", 4:"abril", 5:"mayo", 6:"junio", 7:"julio", 8:"agosto", 9:"septiembre", 10:"octubre", 11:"noviembre", 12:"diciembre"}

with st.form("form_pdc"):
    st.subheader("Datos referenciales")
    c1, c2, c3 = st.columns(3)
    with c1:
        nro_plan = st.number_input("Nº de Plan Curricular", min_value=1, value=4, step=1)
        unidad = st.text_input("Unidad Educativa", value="“Germán Busch”")
        nivel = st.selectbox("Nivel", options=["PRIMARIA", "SECUNDARIA"])
        area = st.text_input("Área (Ej: Técnica Tecnológica)", value="Técnica Tecnológica")
    with c2:
        docente = st.text_input("Maestra/o", value="Verónica Blanca Colmena Quispe")
        ano = st.text_input("Año de Escolaridad", value="Quinto A, B y C.")
        trimestre = st.selectbox("Trimestre", options=["PRIMERO", "SEGUNDO", "TERCERO"], index=1)
        semanas = st.selectbox("Semanas de duración", options=[1,2,3,4], index=3)
    with c3:
        tema = st.text_input("El Tema a avanzar", placeholder="Ej: Las actividades económicas...")
        st.write("Rango de Fechas")
        f1, f2 = st.columns(2)
        with f1:
            fecha_inicio = st.date_input("Desde")
        with f2:
            fecha_fin = st.date_input("Hasta")
            
    puntos_clave = st.text_area("Puntos clave o enfoques específicos (Opcional)", placeholder="Ej: Mencionar billetes de alasitas, hacer énfasis en la honestidad del cambio, etc.")
            
    submitted = st.form_submit_button("Generar Plan en Word")

if submitted:
    if not tema or tema.strip() == "":
        st.error("Por favor ingrese el campo 'El Tema a avanzar'.")
    else:
        mes_ini = MESES_ESPANOL.get(fecha_inicio.month, "enero")
        mes_fin = MESES_ESPANOL.get(fecha_fin.month, "enero")
        cadena_fechas = f"Del: ......................... {fecha_inicio.day} de {mes_ini} ......................... Al: ......... {fecha_fin.day} de {mes_fin} ......... del {fecha_fin.year} ............................................"

        datos_form = {
            "nro_plan": nro_plan,
            "unidad": unidad,
            "nivel": nivel,
            "ano": ano,
            "docente": docente,
            "area": area,
            "trimestre": trimestre,
            "fechas": cadena_fechas,
        }

        try:
            with st.spinner("Generando plan con textos ricos y detallados (exclusivo para Ser, Saber y Hacer)..."):
                client = get_groq_client()
                prompt = build_prompt(tema=tema, semanas=semanas, nivel=nivel, area=area, puntos_clave=puntos_clave)

                response = client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "Eres un pedagogo experto y detallista. Tu prioridad ABSOLUTA es generar textos largos, profundos y muy descriptivos. Responde ÚNICAMENTE con formato JSON."},
                        {"role": "user", "content": prompt}
                    ],
                    model="llama-3.1-70b-versatile",
                    temperature=0.4, 
                )

                text = response.choices[0].message.content
                plan_json = extract_json_from_text(text)
                docx_bio = generar_docx(datos_form, plan_json)

                st.success("¡Plan generado con contenido detallado, extenso y completo!")
                now_name = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"PDC_{nivel}_{tema[:10]}_{now_name}.docx"
                st.download_button(
                    label="📥 Descargar Plan Curricular",
                    data=docx_bio.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Ocurrió un error: {e}")
            with st.expander("Detalles del error"):
                st.exception(e)